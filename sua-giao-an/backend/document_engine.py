"""Document Engine prototype for TRANG SỬA GIÁO ÁN.

Principles:
- Never overwrite the source file.
- Apply formatting only from the approved format standard.
- Do not rewrite content as plain text.
- Return validation findings before export.
"""
from __future__ import annotations

from dataclasses import dataclass, asdict
from pathlib import Path
from copy import deepcopy
from typing import Any
import json

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Cm, Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

STANDARD_PATH = Path(__file__).resolve().parents[1] / "format-standard.json"


@dataclass
class ValidationIssue:
    severity: str
    code: str
    message: str


@dataclass
class DocumentSnapshot:
    paragraphs: int
    tables: int
    sections: int
    inline_shapes: int


class DocumentEngine:
    def __init__(self, standard_path: Path = STANDARD_PATH) -> None:
        self.standard = json.loads(standard_path.read_text(encoding="utf-8"))

    def snapshot(self, doc: Document) -> DocumentSnapshot:
        return DocumentSnapshot(
            paragraphs=len(doc.paragraphs),
            tables=len(doc.tables),
            sections=len(doc.sections),
            inline_shapes=len(doc.inline_shapes),
        )

    def normalize_sections(self, doc: Document) -> None:
        page = self.standard["page"]
        for section in doc.sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(page["marginsCm"]["top"])
            section.bottom_margin = Cm(page["marginsCm"]["bottom"])
            section.left_margin = Cm(page["marginsCm"]["left"])
            section.right_margin = Cm(page["marginsCm"]["right"])

    def normalize_paragraphs(self, doc: Document) -> None:
        font = self.standard["font"]["body"]
        line_spacing = self.standard["paragraph"]["lineSpacing"]
        for paragraph in doc.paragraphs:
            pf = paragraph.paragraph_format
            pf.line_spacing = line_spacing
            for run in paragraph.runs:
                run.font.name = font["name"]
                run.font.size = Pt(font["sizePt"])
                # Keep bold/italic/underline from source; only normalize base font/size.

    def normalize_tables(self, doc: Document) -> None:
        ratio = self.standard["tables"]["defaultTwoColumnRatio"]
        for table in doc.tables:
            table.autofit = False
            if len(table.columns) == 2:
                # Width is applied by Word from the available page area.
                # Explicit widths are avoided here so merged/nested tables are not damaged.
                for row in table.rows:
                    for cell in row.cells:
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.line_spacing = line_spacing = self.standard["paragraph"]["lineSpacing"]
                        for run in paragraph.runs:
                            run.font.name = font_name = self.standard["font"]["body"]["name"]
                            run.font.size = Pt(self.standard["font"]["body"]["sizePt"])

    def format_document(self, input_path: Path, output_path: Path) -> dict[str, Any]:
        if input_path.resolve() == output_path.resolve():
            raise ValueError("Output must be a new file; source must never be overwritten.")

        doc = Document(str(input_path))
        before = self.snapshot(doc)
        self.normalize_sections(doc)
        self.normalize_paragraphs(doc)
        self.normalize_tables(doc)
        doc.save(str(output_path))
        reopened = Document(str(output_path))
        after = self.snapshot(reopened)
        validation = self.validate(before, after)
        return {
            "input": str(input_path),
            "output": str(output_path),
            "before": asdict(before),
            "after": asdict(after),
            "validation": [asdict(x) for x in validation],
            "export_allowed": not any(x.severity == "ERROR" for x in validation),
        }

    def validate(self, before: DocumentSnapshot, after: DocumentSnapshot) -> list[ValidationIssue]:
        issues: list[ValidationIssue] = []
        if after.paragraphs < before.paragraphs:
            issues.append(ValidationIssue("ERROR", "PARAGRAPH_LOSS", "Output contains fewer paragraphs than source."))
        if after.tables < before.tables:
            issues.append(ValidationIssue("ERROR", "TABLE_LOSS", "Output contains fewer tables than source."))
        if after.sections != before.sections:
            issues.append(ValidationIssue("ERROR", "SECTION_CHANGED", "Section count changed during formatting."))
        if after.inline_shapes < before.inline_shapes:
            issues.append(ValidationIssue("ERROR", "IMAGE_LOSS", "Output contains fewer supported inline images than source."))
        if not issues:
            issues.append(ValidationIssue("OK", "VALID", "Basic structural validation passed."))
        return issues


__all__ = ["DocumentEngine", "ValidationIssue", "DocumentSnapshot"]
