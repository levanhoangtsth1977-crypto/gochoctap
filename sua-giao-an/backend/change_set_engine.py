from __future__ import annotations

import json
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
CURRICULUM_PATH = ROOT / "curriculum-adjustments-2026-2027.json"
ADMIN_PATH = ROOT / "administrative-standard-2026-2027.json"
ZONE_PATH = ROOT / "special-administrative-zones-2026-2027.json"

@dataclass
class ChangeSet:
    id: str
    engine: str
    status: str
    severity: str
    location: str
    old: str
    new: str
    reason: str
    confidence: float
    source_ref: str

class ChangeSetEngine:
    """Deterministic DOCX inspection. Proposes only; never mutates the source during inspect."""

    def __init__(self) -> None:
        self.curriculum = json.loads(CURRICULUM_PATH.read_text(encoding="utf-8"))
        self.admin = json.loads(ADMIN_PATH.read_text(encoding="utf-8"))
        self.zones = json.loads(ZONE_PATH.read_text(encoding="utf-8")) if ZONE_PATH.exists() else {}
        self.rules: list[dict[str, Any]] = []
        for subject, rows in self.curriculum.get("subjects", {}).items():
            for row in rows:
                item = dict(row)
                item["subject"] = subject
                self.rules.append(item)

    @staticmethod
    def _norm(text: str) -> str:
        text = (text or "").replace("–", "-").replace("—", "-")
        return re.sub(r"\s+", " ", text).strip().casefold()

    @staticmethod
    def _compact(text: str) -> str:
        return re.sub(r"[^0-9a-zA-ZÀ-ỹĐđ]+", " ", text or "").strip().casefold()

    def _rule_context_ok(self, rule: dict[str, Any], metadata: dict[str, str]) -> bool:
        supplied_subject = self._norm(metadata.get("subject", ""))
        supplied_lesson = self._norm(metadata.get("lesson", ""))
        rule_subject = self._norm(rule.get("subject", ""))
        aliases = {
            "đạo đức": "dao_duc", "lịch sử - địa lí": "lich_su_dia_li", "lịch sử–địa lí": "lich_su_dia_li",
            "lịch sử - địa lý": "lich_su_dia_li", "lịch sử-địa lý": "lich_su_dia_li",
            "toán": "toan", "tiếng việt": "tieng_viet", "khoa học": "khoa_hoc",
        }
        if supplied_subject:
            mapped = aliases.get(supplied_subject, supplied_subject)
            if mapped != rule_subject:
                return False
        rule_lesson = self._norm(rule.get("lesson", ""))
        if supplied_lesson and rule_lesson and rule_lesson not in supplied_lesson and supplied_lesson not in rule_lesson:
            return False
        return True

    def _find_occurrences(self, full_text: str, rule: dict[str, Any]) -> list[tuple[str, str]]:
        """Find exact or natural-language variants of the legacy phrase.

        Returns (matched_text, reason). This fixes cases such as
        'Cột cờ Lũng Cú ở tỉnh Hà Giang' vs rule 'Cột cờ Lũng Cú (Hà Giang)'.
        """
        old = str(rule.get("old", "")).strip()
        if not old:
            return []
        hay = self._norm(full_text)
        old_n = self._norm(old)
        matches: list[tuple[str, str]] = []
        if old_n in hay:
            for m in re.finditer(re.escape(old_n), hay, flags=re.I):
                matches.append((m.group(0), "exact"))
            return matches

        # Build semantic anchors from meaningful locality tokens.
        # We deliberately use the stable landmark/topic words, while avoiding a blind global replace.
        def tokens(s: str) -> list[str]:
            return [x for x in re.split(r"[^0-9A-Za-zÀ-ỹĐđ]+", self._norm(s)) if len(x) >= 2]
        ot = tokens(old)
        if len(ot) < 2:
            return []
        landmark = [x for x in ot if x not in {"tỉnh", "thành", "phố", "huyện", "xã", "đảo", "quần", "cột", "cờ", "một", "góc"}]
        if not landmark:
            landmark = ot
        # For each unique landmark family, accept flexible linking words and legacy province names.
        if len(landmark) >= 2:
            core = landmark[:6]
            pattern = r"(?i)(?:" + r"\s+".join(map(re.escape, core)) + r")(?:\s+(?:ở|tại|thuộc|nằm|trong))?(?:\s+(?:tỉnh|thành phố|huyện|xã))?\s+[^\n,.;:()]{2,80}"
            for m in re.finditer(pattern, self._norm(full_text)):
                matched = m.group(0).strip()
                # Verify the matched span carries the old locality token where applicable.
                old_place = ot[-1]
                if old_place in matched or any(p in matched for p in ("hà giang", "bà rịa - vũng tàu", "quảng bình", "kiên giang", "bắc kạn", "bắc giang", "ninh thuận", "quảng nam", "bình định", "nam định", "hải dương", "thái bình", "yên bái", "bạc liêu", "nghệ an", "cà mau")):
                    matches.append((matched, "variant"))
        # Simpler direct variant: landmark + a legacy province/city enclosed in parentheses.
        if len(ot) >= 2:
            label = " ".join(landmark[:6])
            pat = rf"(?i){r'\s+'.join(map(re.escape, label.split()))}\s*\([^)]*\)"
            for m in re.finditer(pat, self._norm(full_text)):
                if m.group(0) not in {x[0] for x in matches}:
                    matches.append((m.group(0), "variant"))
        # De-duplicate.
        seen = set(); out=[]
        for x in matches:
            if x[0] not in seen:
                seen.add(x[0]); out.append(x)
        return out

    def _curriculum_changes(self, full_text: str, metadata: dict[str, str]) -> list[ChangeSet]:
        out: list[ChangeSet] = []
        for row in self.rules:
            if not self._rule_context_ok(row, metadata):
                continue
            found = self._find_occurrences(full_text, row)
            if not found:
                continue
            action = row.get("action", "review")
            safe = action in {"replace", "replace_contextual", "remove_legacy_district", "remove_legacy_province"}
            for idx, (matched, mode) in enumerate(found, 1):
                status = "PROPOSED" if safe else "REVIEW_REQUIRED"
                severity = "WARNING" if not safe else "INFO"
                # Preserve the canonical source phrase in the reason while carrying the actual matched text as anchor.
                out.append(ChangeSet(
                    id=f"{row['id']}#{idx}" if len(found) > 1 else row["id"],
                    engine="Curriculum Adjustment Engine",
                    status=status,
                    severity=severity,
                    location=f"{row.get('lesson','')}; {row.get('location','')}".strip("; "),
                    old=matched,
                    new=row.get("new", ""),
                    reason=(row.get("note", "Đối chiếu danh mục điều chỉnh SGK Kết nối tri thức 2026–2027.")
                            + (" Phát hiện theo biến thể diễn đạt tự nhiên trong văn bản." if mode == "variant" else "")),
                    confidence=0.985 if mode == "exact" else 0.97,
                    source_ref=f"curriculum-adjustments-2026-2027.json#{row['id']}",
                ))
        return out

    def _administrative_changes(self, full_text: str) -> list[ChangeSet]:
        out: list[ChangeSet] = []
        checks = [
            ("ADMIN-63", r"\b63\s+(?:tỉnh|tỉnh,)?\s*(?:thành phố(?:\s+trực\s+thuộc\s+Trung\s+ương)?)?", "34 đơn vị hành chính cấp tỉnh, gồm 27 tỉnh và 7 thành phố trực thuộc Trung ương", "Dữ liệu cũ trước chuẩn 2026–2027."),
            ("ADMIN-28-6", r"\b28\s+tỉnh\s*(?:và|\+)\s*6\s+thành phố", "27 tỉnh + 7 thành phố trực thuộc Trung ương", "Cơ cấu cũ; cần cập nhật theo chuẩn 2026–2027."),
            ("ADMIN-5", r"\b5\s+thành phố\s+trực\s+thuộc\s+Trung\s+ương", "7 thành phố trực thuộc Trung ương", "Cơ cấu cũ; cần cập nhật theo chuẩn 2026–2027."),
            ("ADMIN-63-PROSE", r"\b63\s+tỉnh\s*,?\s*thành\s+phố\s+trực\s+thuộc\s+trung\s+ương\b", "34 đơn vị hành chính cấp tỉnh, gồm 27 tỉnh và 7 thành phố trực thuộc Trung ương", "Cơ cấu cũ; cần cập nhật theo chuẩn 2026–2027."),
        ]
        hay = self._norm(full_text)
        for cid, pattern, new, reason in checks:
            for idx, m in enumerate(re.finditer(pattern, hay, flags=re.I), 1):
                out.append(ChangeSet(f"{cid}#{idx}" if len(re.findall(pattern, hay, flags=re.I)) > 1 else cid, "Administrative Data Engine", "PROPOSED", "WARNING", "Toàn văn tài liệu", m.group(0), new, reason, 0.995, f"administrative-standard-2026-2027.json#{cid}"))

        # Specific province/city legacy names: review rather than blind global replacement.
        province_moves = {
            "Hà Giang": "Tuyên Quang", "Bà Rịa – Vũng Tàu": "Thành phố Hồ Chí Minh", "Bà Rịa - Vũng Tàu": "Thành phố Hồ Chí Minh",
            "Quảng Bình": "Quảng Trị", "Kiên Giang": "An Giang", "Bắc Kạn": "Thái Nguyên", "Bắc Giang": "Bắc Ninh",
            "Ninh Thuận": "Khánh Hòa", "Quảng Nam": "Đà Nẵng", "Bình Định": "Gia Lai", "Nam Định": "Ninh Bình",
            "Hải Dương": "Hải Phòng", "Thái Bình": "Hưng Yên", "Yên Bái": "Lào Cai", "Bạc Liêu": "Cà Mau",
        }
        for old, new in province_moves.items():
            # Only emit contextual review proposals when the legacy locality is explicitly present.
            pats = [rf"\b(?:tỉnh\s+)?{re.escape(old)}\b"]
            for pat in pats:
                matches = list(re.finditer(pat, hay, flags=re.I))
                for idx, m in enumerate(matches, 1):
                    out.append(ChangeSet(
                        f"ADMIN-PLACE-{self._compact(old).replace(' ', '-')}#{idx}" if len(matches) > 1 else f"ADMIN-PLACE-{self._compact(old).replace(' ', '-')}",
                        "Administrative Data Engine", "REVIEW_REQUIRED", "WARNING", "Toàn văn tài liệu", m.group(0), new,
                        "Phát hiện địa danh cấp tỉnh có thể là tên cũ; phải đối chiếu đúng bài/ngữ cảnh trước khi sửa.",
                        0.94, f"administrative-standard-2026-2027.json#place:{old}"))

        m = re.search(r"\btỉnh\s+Đồng\s+Nai\b", hay, flags=re.I)
        if m:
            out.append(ChangeSet("ADMIN-DONGNAI", "Administrative Data Engine", "REVIEW_REQUIRED", "WARNING", "Toàn văn tài liệu", m.group(0), "Thành phố Đồng Nai", "Đối chiếu ngữ cảnh và thời điểm văn bản trước khi sửa.", 0.98, "administrative-standard-2026-2027.json#dongNai"))
        return out

    def _zone_changes(self, full_text: str) -> list[ChangeSet]:
        out: list[ChangeSet] = []
        rows = self.zones.get("zones", []) if isinstance(self.zones, dict) else []
        hay = self._norm(full_text)
        for row in rows:
            name = row.get("name") or row.get("zone") or ""
            if not name:
                continue
            pattern = rf"\bhuyện\s+{re.escape(name)}\b"
            for idx, m in enumerate(re.finditer(pattern, hay, flags=re.I), 1):
                out.append(ChangeSet(
                    f"ZONE-{self._compact(name).replace(' ', '-').upper()}#{idx}" if len(re.findall(pattern, hay, flags=re.I)) > 1 else f"ZONE-{self._compact(name).replace(' ', '-').upper()}",
                    "Special Administrative Zone Engine", "REVIEW_REQUIRED", "WARNING", "Toàn văn tài liệu", m.group(0), f"đặc khu {name}",
                    f"Kiểm tra cách gọi đơn vị hành chính hiện hành của đặc khu {name}; không sửa tự động nếu chưa xác định ngữ cảnh.",
                    0.96, f"special-administrative-zones-2026-2027.json#{name}"))
        return out

    def inspect(self, path: Path, metadata: dict[str, str] | None = None) -> dict[str, Any]:
        metadata = metadata or {}
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        full_text = "\n".join(x for x in parts if x.strip())
        changes = self._curriculum_changes(full_text, metadata) + self._administrative_changes(full_text) + self._zone_changes(full_text)
        return {
            "document": path.name,
            "snapshot": {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables), "sections": len(doc.sections), "inlineShapes": len(doc.inline_shapes)},
            "context": metadata,
            "changes": [asdict(x) for x in changes],
            "changeCount": len(changes),
            "exportAllowed": False,
            "message": "Đã quét toàn văn DOCX, kể cả nội dung bảng, và tạo Change Set; chưa sửa file gốc.",
        }

__all__ = ["ChangeSetEngine", "ChangeSet"]
