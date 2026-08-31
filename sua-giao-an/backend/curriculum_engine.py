from __future__ import annotations

from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
import json
import re
import zipfile
import io


@dataclass
class ChangeSet:
    id: str
    engine: str
    action: str
    status: str
    confidence: str
    location: str
    old_text: str
    new_text: str
    reason: str
    evidence_id: str = ""


class KnowledgeBase:
    """Deterministic lookup layer. AI is not used as the authority for facts."""

    def __init__(self, root: Path):
        self.root = root
        self.curriculum = self._load_json("curriculum-adjustments-2026-2027.json")
        self.admin = self._load_json("administrative-standard-2026-2027.json")
        self.zones = self._load_json("special-administrative-zones-2026-2027.json")

    def _load_json(self, name: str) -> Dict[str, Any]:
        path = self.root / name
        return json.loads(path.read_text(encoding="utf-8"))

    @staticmethod
    def norm(value: str) -> str:
        return re.sub(r"\s+", " ", value.casefold()).strip()

    def curriculum_rules(self, subject: str) -> List[Dict[str, Any]]:
        aliases = {
            "Đạo đức": "DAO_DUC",
            "DAO_DUC": "DAO_DUC",
            "Lịch sử - Địa lí": "LICH_SU_DIA_LI",
            "Lịch sử–Địa lí": "LICH_SU_DIA_LI",
            "LICH_SU_DIA_LI": "LICH_SU_DIA_LI",
            "Toán": "TOAN",
            "TOAN": "TOAN",
            "Tiếng Việt": "TIENG_VIET",
            "TIENG_VIET": "TIENG_VIET",
        }
        return self.curriculum.get("subjects", {}).get(aliases.get(subject, subject), [])

    def administrative_unit(self, name: str) -> Optional[str]:
        n = self.norm(name)
        for x in self.admin.get("municipalities", []):
            if self.norm(x) == n:
                return "thành phố trực thuộc Trung ương"
        for x in self.admin.get("provinceList", []):
            if self.norm(x) == n:
                return "tỉnh"
        return None

    def zone(self, name: str) -> Optional[Dict[str, Any]]:
        n = self.norm(name)
        for z in self.zones.get("zones", []):
            if self.norm(z.get("name", "")) == n:
                return z
        return None


def _contains(text: str, needle: str) -> bool:
    return KnowledgeBase.norm(needle) in KnowledgeBase.norm(text)


def _replace_cs(rule: Dict[str, Any], location: str, idx: int) -> Optional[ChangeSet]:
    if not rule.get("old") or not rule.get("new"):
        return None
    return ChangeSet(
        id=f"CS-{idx:04d}",
        engine="Curriculum Adjustment Engine",
        action=rule.get("action", "review"),
        status="PROPOSED",
        confidence="HIGH" if rule.get("action") in {"replace", "remove_legacy_province", "remove_legacy_district"} else "MEDIUM",
        location=location,
        old_text=rule["old"],
        new_text=rule["new"],
        reason="Khớp bản ghi điều chỉnh SGK theo môn/bài/trang/mục.",
        evidence_id=rule.get("id", ""),
    )


def build_change_sets(text: str, metadata: Dict[str, str], kb: KnowledgeBase) -> List[ChangeSet]:
    """Create proposals only; never mutates the document."""
    out: List[ChangeSet] = []
    subject = metadata.get("subject", "")
    lesson = KnowledgeBase.norm(metadata.get("lesson", ""))
    location = metadata.get("location", "Toàn văn")
    idx = 1

    # 1) Exact curriculum adjustment records: require textual match + lesson/location context.
    for rule in kb.curriculum_rules(subject):
        rule_lesson = KnowledgeBase.norm(rule.get("lesson", ""))
        rule_location = KnowledgeBase.norm(rule.get("location", ""))
        context_ok = not rule_lesson or rule_lesson in lesson or lesson in rule_lesson
        if rule_location:
            context_ok = context_ok and (rule_location in KnowledgeBase.norm(location) or rule_location in KnowledgeBase.norm(text))
        if context_ok and _contains(text, rule.get("old", "")):
            cs = _replace_cs(rule, location, idx)
            if cs:
                out.append(cs); idx += 1

    # 2) Administrative legacy patterns: flag, do not invent replacements.
    admin_legacy = kb.admin.get("legacyPatterns", [])
    for item in admin_legacy:
        pattern = item.get("pattern", "")
        if pattern and _contains(text, pattern):
            out.append(ChangeSet(
                id=f"CS-{idx:04d}", engine="Administrative Data Engine",
                action=item.get("action", "flag_outdated"), status="REVIEW_REQUIRED",
                confidence="HIGH", location=location, old_text=pattern,
                new_text="", reason="Phát hiện mẫu dữ liệu hành chính cũ; cần đối chiếu ngữ cảnh trước khi sửa.",
                evidence_id="ADMIN-LEGACY",
            )); idx += 1

    # 3) Special administrative zones: detect old unit prefixes, but only propose when exact zone name exists.
    for z in kb.zones.get("zones", []):
        name = z.get("name", "")
        if not name or not _contains(text, name):
            continue
        for old_prefix in ("huyện ", "Huyện "):
            old = old_prefix + name
            if old in text:
                out.append(ChangeSet(
                    id=f"CS-{idx:04d}", engine="Special Administrative Zone Engine",
                    action="context_check", status="REVIEW_REQUIRED", confidence="MEDIUM",
                    location=location, old_text=old, new_text=f"đặc khu {name}",
                    reason=f"Đối chiếu danh mục đặc khu hiện hành; chỉ sửa khi văn bản đang nói về đơn vị hành chính.",
                    evidence_id=f"ZONE-{name}",
                )); idx += 1
    return out


def analyze_docx_bytes(data: bytes, metadata: Dict[str, str], kb_root: Path) -> Dict[str, Any]:
    """Lightweight structural extraction without rewriting the DOCX."""
    try:
        from docx import Document
    except Exception as exc:
        return {"ok": False, "error": f"python-docx unavailable: {exc}"}

    temp = io.BytesIO(data)
    doc = Document(temp)
    paragraphs = [p.text for p in doc.paragraphs]
    tables = []
    for ti, table in enumerate(doc.tables, start=1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        tables.append({"index": ti, "rows": len(rows), "cols": max((len(r) for r in rows), default=0), "cells": rows})
    text = "\n".join(paragraphs)
    for table in tables:
        for row in table["cells"]:
            text += "\n" + " | ".join(row)

    kb = KnowledgeBase(kb_root)
    changes = build_change_sets(text, metadata, kb)
    return {
        "ok": True,
        "document": {"paragraph_count": len(paragraphs), "table_count": len(tables), "tables": tables},
        "change_sets": [asdict(x) for x in changes],
        "mutated": False,
    }
